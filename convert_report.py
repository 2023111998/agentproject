#!/usr/bin/env python3
"""Convert 课程报告.md to Word .docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

for level in range(1, 4):
    h_style = doc.styles['Heading %d' % level]
    h_font = h_style.font
    h_font.name = 'Times New Roman'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    h_font.bold = True
    h_font.color.rgb = RGBColor(0, 0, 0)
    sizes = {1: 16, 2: 14, 3: 13}
    h_font.size = Pt(sizes.get(level, 12))

def add_para(text, bold=False, size=12, center=False, font_cn='SimSun'):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(8.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold = True
                rr.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri + 1].cells[ci]
            c.text = str(val)
            for pp in c.paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(9)
    doc.add_paragraph()

def add_hr():
    p = doc.add_paragraph()
    p.add_run('_' * 60)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def clean_markdown(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text

# Parse markdown
with open(r'D:\lab\Agent服务工程\campus-assistant-java\课程报告.md', 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
i = 0
in_code = False
code_buf = []
in_ascii = False
ascii_buf = []

while i < len(lines):
    line = lines[i]
    s = line.strip()

    # Code block toggle
    if s.startswith('```'):
        if in_code:
            for cl in code_buf:
                add_code(cl)
            code_buf = []
            in_code = False
        else:
            in_code = True
        i += 1
        continue

    if in_code:
        code_buf.append(s)
        i += 1
        continue

    # ASCII diagram detection
    ascii_chars = set('|.-+=/\\<>[](){}_^~*#@$%&!:;?')
    is_ascii_art = any(c in s for c in '│┌└├┬┴┼─╭╮╰╯┤┘┐┃━┏┓┗┛┣┫┳┻╋') or \
        (len(s) > 30 and all(c in ascii_chars or c.isspace() or c.isalpha() or c.isdigit() for c in s) and s.count('─') > 2)

    if is_ascii_art:
        if not in_ascii:
            in_ascii = True
            ascii_buf = []
        ascii_buf.append(line)
        i += 1
        continue
    elif in_ascii:
        for al in ascii_buf:
            add_code(al.rstrip())
        ascii_buf = []
        in_ascii = False

    # Headings
    heading_match = re.match(r'^(#{1,3})\s+(.+)', s)
    if heading_match:
        lvl = len(heading_match.group(1))
        txt = clean_markdown(heading_match.group(2))
        doc.add_heading(txt, level=lvl)
        i += 1
        continue

    # HR
    if s == '---':
        add_hr()
        i += 1
        continue

    # Table detection
    if s.startswith('|') and s.endswith('|'):
        # Check if separator follows
        if i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip()):
            headers = [c.strip() for c in s.split('|')[1:-1]]
            data = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith('|') and lines[j].strip().endswith('|'):
                data.append([c.strip() for c in lines[j].strip().split('|')[1:-1]])
                j += 1
            add_table(headers, data)
            i = j
            continue
        else:
            i += 1
            continue

    # Regular text
    if s:
        clean = clean_markdown(s)
        add_para(clean, size=12)
    else:
        doc.add_paragraph()

    i += 1

# Flush any remaining buffers
if in_ascii:
    for al in ascii_buf:
        add_code(al.rstrip())

output = r'D:\lab\Agent服务工程\campus-assistant-java\课程报告.docx'
doc.save(output)
print('Report saved: ' + str(len(content)) + ' chars -> ' + output)
